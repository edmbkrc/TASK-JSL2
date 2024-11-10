import os
import tempfile
import streamlit as st
from langchain.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
import textract
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.schema.document import Document


def llama_rag_main():
    # Prompt Template
    prompt_template = """
    Given the following context, answer the question as accurately and comprehensively as possible.
    Combine relevant information from all provided chunks if necessary. 
    If the answer cannot be derived from the context, respond only with: "Sorry, I can only answer questions based on the provided documents.
    Context:
    {context}

    Question:
    {input}

    Answer:
    """

    # Function to read DOCX files
    def read_docx(file):
        doc = DocxDocument(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to read DOC files
    def read_doc(file):
        return textract.process(file, encoding="utf-8").decode("utf-8")

    # Function to update vector store dynamically
    def update_vector_store(uploaded_files, query=None):
        for uploaded_file in uploaded_files:
            if uploaded_file.name in st.session_state.processed_files:
                st.write(f"Processed file: {uploaded_file.name}")
                continue

            file_extension = uploaded_file.name.split(".")[-1].lower()
            all_docs = []

            if file_extension == "pdf":
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                    temp_file.write(uploaded_file.read())
                    temp_path = temp_file.name
                pdf_loader = PyPDFLoader(temp_path)
                pdf_docs = pdf_loader.load()
                all_docs.extend([Document(page_content=doc.page_content, metadata=doc.metadata) for doc in pdf_docs])
                os.remove(temp_path)

            elif file_extension == "docx":
                text = read_docx(uploaded_file)
                if text.strip():
                    all_docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))

            elif file_extension == "doc":
                with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                    temp_file.write(uploaded_file.read())
                    temp_path = temp_file.name
                text = read_doc(temp_path)
                if text.strip():
                    all_docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
                os.remove(temp_path)

            st.session_state.processed_files.add(uploaded_file.name)

            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
            new_chunks = text_splitter.split_documents(all_docs)

            if not new_chunks:
                st.warning(f"No valid content found in {uploaded_file.name}. Skipping.")
                continue
            
            # Apply filtering if a query is provided
            if query:
                def filter_relevant_chunks(query, chunks):
                    query_keywords = query.lower().split()
                    filtered_chunks = [
                        chunk for chunk in chunks 
                        if any(keyword in chunk.page_content.lower() for keyword in query_keywords)
                    ]
                    return filtered_chunks if filtered_chunks else chunks  # Default to all if none match

                filtered_chunks = filter_relevant_chunks(query, new_chunks)
            else:
                filtered_chunks = new_chunks

            # Convert chunks into embeddings
            texts = [chunk.page_content for chunk in filtered_chunks]
            metadatas = [chunk.metadata for chunk in filtered_chunks]
            embeddings = st.session_state.embeddings.embed_documents(texts)

            # Update or create the vectorstore
            if st.session_state.vectors is not None:
                st.session_state.vectors.add_texts(texts=texts, embeddings=embeddings, metadatas=metadatas)
                st.write(f"Vectorstore updated with new chunks from {uploaded_file.name}!")
            else:
                st.session_state.vectors = FAISS.from_texts(texts=texts, embedding=st.session_state.embeddings, metadatas=metadatas)
                st.write(f"Vectorstore created with chunks from {uploaded_file.name}!")

    # Initialize Session State
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = set()
    if "embeddings" not in st.session_state:
        st.session_state.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L12-v2")
    if "vectors" not in st.session_state:
        st.session_state.vectors = None
    if "cache" not in st.session_state:
        st.session_state.cache = {}
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Streamlit UI
    st.title("RAG Chatbot with Llama")

    # File uploader
    uploaded_files = st.file_uploader(
        "Upload your documents (PDF, DOC, DOCX)", type=["pdf", "doc", "docx"], accept_multiple_files=True
    )

    if uploaded_files:
        with st.spinner("Processing files..."):
            update_vector_store(uploaded_files)

    # Chat interface
    st.subheader("Chat with your documents")
    user_input = st.text_input("You:", key="input")

    if user_input:
        # Check if the query exists in the cache
        if user_input in st.session_state.cache:
            cached_response = st.session_state.cache[user_input]
            st.write("**Cached Response:** This response is retrieved from the cache.")
            st.write(f"**You:** {user_input}")
            st.write(f"**Bot:** {cached_response['answer']}")
            with st.expander("Retrieved Chunks"):
                for i, doc in enumerate(cached_response["retrieved_docs"], start=1):
                    page_number = doc.metadata.get("page", "Unknown")
                    st.write(f"**Chunk {i} (Page {str(page_number) if isinstance(page_number, int) else page_number}):**")
                    st.write(doc.page_content)
                    st.write("---")
        else:
            # Generate embeddings for new queries
            retriever = st.session_state.vectors.as_retriever(search_kwargs={"k": 5, "distance_threshold": 0.7})
            retrieved_docs = retriever.get_relevant_documents(user_input)

            if retrieved_docs:
                context = "\n".join([doc.page_content for doc in retrieved_docs])
                formatted_prompt = prompt_template.format(context=context, input=user_input)

                llm = Ollama(
                    model="llama3.2:1b",
                    host="http://localhost:11434",
                    callback_manager=CallbackManager([StreamingStdOutCallbackHandler()])
                )
                response = llm(formatted_prompt)
            else:
                response = "Sorry, I can only answer questions based on the provided documents. No relevant information was found."

            st.session_state.cache[user_input] = {"answer": response, "retrieved_docs": retrieved_docs}
            st.session_state.messages.append((user_input, response))

            st.write(f"**You:** {user_input}")
            st.write(f"**Bot:** {response}")

            with st.expander("Retrieved Chunks"):
                for i, doc in enumerate(retrieved_docs, start=1):
                    source = doc.metadata.get("source", "Unknown")
                    page = doc.metadata.get("page", "Unknown")
                    st.write(f"**Chunk {i} (Source: {source},Page {str(page) if isinstance(page, int) else page}):**")
                    st.write(doc.page_content)
                    st.write("---")

    with st.expander("Chat History"):
        st.write("### Chat History")
        for user_msg, bot_msg in st.session_state.messages:
            st.write(f"You: {user_msg}")
            st.write(f"Bot: {bot_msg}")
